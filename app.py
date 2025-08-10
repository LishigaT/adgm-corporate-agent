# app.py
import streamlit as st
from docx import Document
from io import BytesIO
import google.generativeai as genai
import json
import os
import glob
import re

from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

# -------------------------
# Config & helpers
# -------------------------
st.set_page_config(page_title="ADGM Corporate Agent – Compliance Checker", layout="wide")
st.title("ADGM Corporate Agent – Compliance & RAG-enabled Reviewer")
st.write("Upload `.docx` files. The app will check required ADGM documents, run RAG-backed checks using Gemini, and produce an annotated docx + JSON report.")

# --- Load Gemini API Key from Streamlit secrets ---
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.warning("Gemini API key not found in Streamlit secrets. You must add `GEMINI_API_KEY` in Streamlit secrets to enable AI checks.")
    GEMINI_API_KEY = None

# --- Required documents per process (extendable) ---
REQUIRED_DOCS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ],
    # add more processes if needed
}

# -------------------------
# Utility functions
# -------------------------
def extract_text_from_docx(file_obj):
    """Return full text string from a .docx file-like object."""
    file_obj.seek(0)
    doc = Document(file_obj)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def detect_process(filenames):
    """Basic heuristic to detect process type based on file names (extend as needed)."""
    joined = " ".join(filenames).lower()
    if "articles" in joined or "memorandum" in joined or "incorporation" in joined or "aoa" in joined or "moa" in joined:
        return "Company Incorporation"
    return "Unknown"

def load_reference_texts(ref_dir="references"):
    """
    Load textual ADGM reference files from repo `references/` directory.
    Each file should be plain text (.txt). Returns list of (filename, text).
    """
    refs = []
    if not os.path.exists(ref_dir):
        return refs
    for path in glob.glob(os.path.join(ref_dir, "*.txt")):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                text = fh.read()
            refs.append((os.path.basename(path), text))
        except Exception:
            continue
    return refs

def chunk_text(text, chunk_size=400, overlap=50):
    """
    Simple chunker that splits text into chunks of approx chunk_size words with overlap.
    """
    words = text.split()
    chunks = []
    i = 0
    n = len(words)
    while i < n:
        chunk = words[i:i+chunk_size]
        chunks.append(" ".join(chunk))
        i += chunk_size - overlap
    return chunks

def build_rag_context(document_text, refs, top_k=3):
    """
    Simple TF-IDF retriever:
    - splits reference files into chunks
    - computes TF-IDF similarity between document_text and all chunks
    - returns top_k chunks as the "relevant ADGM references" to include in prompt
    """
    # collect chunks
    chunk_sources = []
    chunk_texts = []
    for fname, text in refs:
        for idx, chunk in enumerate(chunk_text(text, chunk_size=200, overlap=40)):
            chunk_sources.append(f"{fname}::chunk{idx+1}")
            chunk_texts.append(chunk)

    if not chunk_texts:
        return ""

    # TF-IDF vectorizer
    vectorizer = TfidfVectorizer(stop_words="english", max_features=20000)
    # Fit on chunks + the query (document_text)
    try:
        docs = chunk_texts + [document_text]
        X = vectorizer.fit_transform(docs)
        # last vector is query
        query_vec = X[-1]
        chunk_vecs = X[:-1]
        # compute cosine similarities
        sims = (chunk_vecs @ query_vec.T).toarray().ravel()
        top_idx = np.argsort(sims)[-top_k:][::-1]
        top_chunks = [(chunk_sources[i], chunk_texts[i], float(sims[i])) for i in top_idx if sims[i] > 0]
        # format
        formatted = []
        for src, txt, score in top_chunks:
            formatted.append(f"--- Source: {src} (sim={score:.3f}) ---\n{txt}")
        return "\n\n".join(formatted)
    except Exception:
        return ""

def normalize_snippet(sn):
    return re.sub(r"\s+", " ", sn.strip()).lower()

def find_paragraph_indices_containing(doc: Document, snippet: str):
    """
    Return list of paragraph indices that contain snippet (approx match).
    """
    snippet_norm = normalize_snippet(snippet)
    matches = []
    for idx, para in enumerate(doc.paragraphs):
        para_norm = normalize_snippet(para.text)
        if not para_norm:
            continue
        if snippet_norm in para_norm:
            matches.append(idx)
    return matches

def annotate_docx_bytes(original_bytesio, issues):
    """
    Creates an annotated version of the DOCX by adding a clearly-labeled inline comment
    run at the end of paragraphs that match issue snippets (or keywords).
    Returns BytesIO of annotated docx.
    """
    original_bytesio.seek(0)
    doc = Document(original_bytesio)

    for issue in issues:
        doc_name = issue.get("document", "")
        snippet = issue.get("snippet", "") or issue.get("section", "") or issue.get("text", "")
        suggestion = issue.get("suggestion", "")
        severity = issue.get("severity", "")

        if not snippet:
            # fallback: use the 'issue' text as keyword
            snippet = issue.get("issue", "")

        if not snippet:
            continue

        # try to find paragraph(s) matching snippet
        para_indices = find_paragraph_indices_containing(doc, snippet)

        if not para_indices:
            # attempt fuzzy location: look for first paragraph that contains a keyword from snippet
            words = [w for w in re.findall(r"\w{4,}", snippet)][:6]  # shortlist of words length>=4
            for idx, para in enumerate(doc.paragraphs):
                para_norm = normalize_snippet(para.text)
                if any(w.lower() in para_norm for w in words):
                    para_indices.append(idx)
                    break

        # annotate found paragraphs
        for pi in para_indices:
            para = doc.paragraphs[pi]
            # add a new run with a visible annotation (italic + bracketed)
            annotation_text = f" [ADGM_ANNOTATION | Severity: {severity or 'Review'} | Suggestion: {suggestion}]"
            run = para.add_run(annotation_text)
            # style the annotation so it stands out: italic and underline
            run.italic = True
            run.bold = False

    # produce bytes
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# -------------------------
# UI: File upload / processing
# -------------------------
st.markdown("## Upload Documents")
uploaded_files = st.file_uploader("Upload one or more `.docx` files (use sample files if testing)", type=["docx"], accept_multiple_files=True)

# load references for RAG (show count)
refs = load_reference_texts(ref_dir="references")
st.info(f"RAG: loaded {len(refs)} reference file(s) from `references/` (place ADGM .txt references there).")

if uploaded_files:
    # list filenames
    filenames = [f.name for f in uploaded_files]
    st.write("Uploaded files:", filenames)

    # detect process
    process = detect_process(filenames)
    st.write("Detected process:", process)

    # checklist verification
    missing_docs = []
    if process in REQUIRED_DOCS:
        required_list = REQUIRED_DOCS[process]
        missing_docs = [d for d in required_list if d.lower() not in " ".join(filenames).lower()]
        if missing_docs:
            st.error(f" Missing required docs for {process}: {', '.join(missing_docs)} ({len(missing_docs)} missing)")
        else:
            st.success(f" All required documents present for {process} ({len(required_list)} required).")
    else:
        st.info("Unknown/other process - checklist verification skipped.")

    # extract and combine texts for AI + show preview
    combined_texts = {}
    combined_text = ""
    for f in uploaded_files:
        try:
            # need a copy because streamlit file uploader is already a BytesIO-like
            bytes_data = f.getvalue()
            combined = extract_text_from_docx(BytesIO(bytes_data))
            combined_texts[f.name] = {
                "text": combined,
                "bytes": bytes_data
            }
            combined_text += f"\n\n--- {f.name} ---\n\n" + combined
        except Exception as e:
            st.error(f"Error reading {f.name}: {e}")

    st.subheader("Extracted Text (preview)")
    st.text_area("Combined extracted text", combined_text[:10000], height=300)  # limit display

    # RAG build
    rag_context = build_rag_context(combined_text, refs, top_k=3)
    if rag_context:
        st.subheader("RAG: Top relevant ADGM reference passages (used in prompt)")
        st.text_area("RAG context", rag_context[:15000], height=200)
    else:
        st.info("No reference passages found or no `references/` texts available. RAG will be skipped.")

    # run AI
    if GEMINI_API_KEY is None:
        st.warning("Gemini key missing: AI compliance checks disabled. Add `GEMINI_API_KEY` in Streamlit secrets to enable.")
    else:
        if st.button("Run RAG-backed ADGM Compliance Check (Gemini)"):
            with st.spinner("Calling Gemini and analyzing... (this may take a few seconds)"):
                # Prepare prompt
                prompt_parts = [
                    "You are an ADGM Corporate Agent compliance assistant.",
                    "You will be provided with: (A) Relevant ADGM reference passages, and (B) the user's uploaded documents.",
                    "Your task: analyze the documents and return a JSON array named 'issues' where each element has the fields:",
                    " - document: filename where the issue was found",
                    " - snippet: a short excerpt of the text that is the problem (or the exact clause if possible)",
                    " - section: optional (human-readable heading) if available",
                    " - issue: description of the compliance problem",
                    " - severity: one of [High, Medium, Low, Review]",
                    " - suggestion: recommended corrected clause or remediation",
                    "Return ONLY valid JSON. Example:",
                    '{"issues":[{"document":"Articles.docx","snippet":"Clause X ...","section":"Jurisdiction","issue":"Does not reference ADGM","severity":"High","suggestion":"Replace jurisdiction clause with: ..."}]}',
                    ""
                ]
                if rag_context:
                    prompt_parts.append("=== Relevant ADGM reference passages ===")
                    prompt_parts.append(rag_context)
                    prompt_parts.append("=== End of references ===")
                prompt_parts.append("=== Documents to analyze ===")
                # include each file heading + some bytes of text
                for name, info in combined_texts.items():
                    t = info["text"][:8000]  # cap length per file
                    prompt_parts.append(f"--- FILENAME: {name} ---\n")
                    prompt_parts.append(t)
                    prompt_parts.append("\n")
                prompt = "\n".join(prompt_parts)

                # call Gemini
                try:
                    model = genai.GenerativeModel("gemini-1.5-flash")
                    response = model.generate_content(prompt)
                    raw_text = response.text.strip()
                    # attempt JSON parse
                    try:
                        parsed = json.loads(raw_text)
                    except Exception:
                        # Attempt to extract first JSON block from response
                        m = re.search(r"(\{[\s\S]*\})", raw_text)
                        if m:
                            parsed = json.loads(m.group(1))
                        else:
                            # last resort: wrap into analysis field
                            parsed = {"issues": [], "raw": raw_text}

                    issues = parsed.get("issues", [])
                    if not isinstance(issues, list):
                        issues = []

                    st.subheader(" AI Compliance Findings (parsed)")
                    st.write(issues if issues else "No structured issues found. See raw AI output below.")
                    if not issues:
                        st.text_area("Raw AI output", raw_text, height=300)
                except Exception as e:
                    st.error(f"Error calling Gemini: {e}")
                    issues = []

                # Build final report
                final_report = {
                    "process": process,
                    "documents_uploaded": len(filenames),
                    "required_documents": len(REQUIRED_DOCS.get(process, [])) if process in REQUIRED_DOCS else None,
                    "missing_documents": missing_docs,
                    "issues_found": issues
                }

                st.subheader("Downloadable Outputs")
                # JSON report
                st.download_button("⬇ Download JSON Report", json.dumps(final_report, indent=2), file_name="compliance_report.json", mime="application/json")

                # Annotated docx per uploaded file (combine issues by filename)
                # group issues by document
                issues_by_doc = {}
                for it in issues:
                    docname = it.get("document", "")
                    issues_by_doc.setdefault(docname, []).append(it)

                for name, info in combined_texts.items():
                    orig_bytes = BytesIO(info["bytes"])
                    # annotated bytes for the file (pass issues for this filename)
                    doc_issues = issues_by_doc.get(name, [])
                    annotated_bytesio = annotate_docx_bytes(orig_bytes, doc_issues)
                    # original bytes download
                    st.download_button(f"⬇ Download original: {name}", data=info["bytes"], file_name=f"original_{name}", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    # annotated download
                    st.download_button(f"⬇ Download annotated: reviewed_{name}", data=annotated_bytesio.getvalue(), file_name=f"reviewed_{name}", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                st.success("Processing complete. Include the JSON report and the reviewed .docx files as your submission artifacts.")
